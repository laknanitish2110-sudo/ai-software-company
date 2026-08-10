import json
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


GENERATED_DIR = "generated_projects"

ACCENT = RGBColor(99, 91, 255)
DARK = RGBColor(26, 31, 54)
MUTED = RGBColor(82, 95, 127)
SUCCESS = RGBColor(11, 191, 140)
WARNING = RGBColor(245, 166, 35)


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level <= 2 else DARK
    return h


def _para(doc, text, size=11, color=None, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(str(text))
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def _bullet_list(doc, items):
    for item in items:
        bp = doc.add_paragraph(str(item), style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(10)


def _add_structured(doc, content, indent=0):
    if isinstance(content, dict):
        for key, value in content.items():
            if key.startswith("_"):
                continue
            label = key.replace("_", " ").title()
            if isinstance(value, str):
                if len(value) > 200:
                    _heading(doc, label, level=3)
                    _para(doc, value, size=10)
                else:
                    p = doc.add_paragraph()
                    r1 = p.add_run(f"{label}: ")
                    r1.bold = True
                    r1.font.size = Pt(11)
                    r1.font.color.rgb = DARK
                    r2 = p.add_run(value)
                    r2.font.size = Pt(11)
            elif isinstance(value, (int, float, bool)):
                p = doc.add_paragraph()
                r1 = p.add_run(f"{label}: ")
                r1.bold = True
                r1.font.size = Pt(11)
                r2 = p.add_run(str(value))
                r2.font.size = Pt(11)
            elif isinstance(value, list):
                _heading(doc, label, level=3)
                for item in value:
                    if isinstance(item, dict):
                        _add_structured(doc, item, indent + 1)
                    else:
                        bp = doc.add_paragraph(str(item), style="List Bullet")
                        for run in bp.runs:
                            run.font.size = Pt(10)
            elif isinstance(value, dict):
                _heading(doc, label, level=3)
                _add_structured(doc, value, indent + 1)
    elif isinstance(content, list):
        _bullet_list(doc, content)


def _get_output_by_role(outputs, role):
    for o in reversed(outputs):
        if o["role"] == role and o["status"] == "approved":
            return o
    for o in reversed(outputs):
        if o["role"] == role:
            return o
    return None


def _extract(content, *keys):
    if not isinstance(content, dict):
        return None
    for key in keys:
        val = content.get(key)
        if val and isinstance(val, str):
            return val
        if isinstance(val, dict):
            for sv in val.values():
                if isinstance(sv, str) and len(sv) > 20:
                    return sv
    return None


def _extract_list(content, *keys):
    if not isinstance(content, dict):
        return None
    for key in keys:
        val = content.get(key)
        if val and isinstance(val, list):
            return val
    return None


def generate_docx(project_id: str, project: dict, outputs: list, memory: dict) -> str:
    os.makedirs(GENERATED_DIR, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    ceo = _get_output_by_role(outputs, "ceo")
    ba = _get_output_by_role(outputs, "business_analyst")
    researcher = _get_output_by_role(outputs, "researcher")
    architect = _get_output_by_role(outputs, "architect")
    engineer = _get_output_by_role(outputs, "engineer")
    ppt = _get_output_by_role(outputs, "ppt")

    ceo_c = ceo["content"] if ceo and isinstance(ceo.get("content"), dict) else {}
    ba_c = ba["content"] if ba and isinstance(ba.get("content"), dict) else {}
    res_c = researcher["content"] if researcher and isinstance(researcher.get("content"), dict) else {}
    arch_c = architect["content"] if architect and isinstance(architect.get("content"), dict) else {}
    eng_c = engineer["content"] if engineer and isinstance(engineer.get("content"), dict) else {}
    ppt_c = ppt["content"] if ppt and isinstance(ppt.get("content"), dict) else {}

    report = ppt_c.get("report_data", {}) if isinstance(ppt_c.get("report_data"), dict) else {}

    # ═══════════════════════════════════════════
    #  PAGE 1: TITLE PAGE
    # ═══════════════════════════════════════════

    project_name = (
        report.get("title")
        or _extract(ceo_c, "project_name", "title", "name")
        or project.get("problem_statement", "Untitled Project")[:80]
    )

    title = doc.add_heading(project_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = DARK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("AI Software Company — Project Report")
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT

    doc.add_paragraph()

    # ═══════════════════════════════════════════
    #  SECTION 1: WHAT'S THE TITLE OF THE IDEA
    # ═══════════════════════════════════════════

    _heading(doc, "1. Title of the Idea")
    _para(doc, project_name, size=14, bold=True, color=DARK)
    doc.add_paragraph()

    # ═══════════════════════════════════════════
    #  SECTION 2: WHAT IT DOES
    # ═══════════════════════════════════════════

    _heading(doc, "2. What It Does")
    what_it_does = (
        report.get("what_it_does")
        or _extract(ceo_c, "vision", "description", "executive_summary", "overview", "problem_summary")
        or _extract(ba_c, "product_overview", "description", "executive_summary")
        or project.get("problem_statement", "")
    )
    _para(doc, what_it_does)
    doc.add_paragraph()

    # ═══════════════════════════════════════════
    #  SECTION 3: WHAT'S THE PROBLEM
    # ═══════════════════════════════════════════

    _heading(doc, "3. The Problem")
    the_problem = (
        report.get("the_problem")
        or _extract(ceo_c, "problem_analysis", "problem_statement", "problem_summary", "problem", "challenge")
        or _extract(ba_c, "problem_analysis", "problem_statement", "pain_points", "problem")
        or project.get("problem_statement", "")
    )
    _para(doc, the_problem)
    doc.add_paragraph()

    # ═══════════════════════════════════════════
    #  SECTION 4: WHAT IT SOLVES
    # ═══════════════════════════════════════════

    _heading(doc, "4. What It Solves")
    what_it_solves = (
        report.get("what_it_solves")
        or _extract(ceo_c, "recommended_approach", "solution", "value_proposition", "key_outcomes")
        or _extract(ba_c, "solution", "value_proposition", "proposed_solution")
        or _extract(arch_c, "architecture_overview", "solution_overview")
        or "See detailed sections below."
    )
    _para(doc, what_it_solves)
    doc.add_paragraph()

    # ═══════════════════════════════════════════
    #  SECTION 5: FUTURE SCOPE
    # ═══════════════════════════════════════════

    doc.add_page_break()
    _heading(doc, "5. Future Scope")

    future_items = report.get("future_scope")
    if future_items and isinstance(future_items, list):
        _bullet_list(doc, future_items)
    else:
        future_text = (
            _extract(ceo_c, "future_vision", "next_steps", "roadmap")
            or _extract(arch_c, "scalability_strategy", "future", "roadmap")
        )
        future_list = (
            _extract_list(eng_c, "next_steps", "future_work")
            or _extract_list(arch_c, "development_phases")
        )
        if future_list:
            _bullet_list(doc, future_list)
        elif future_text:
            _para(doc, future_text)
        else:
            _para(doc, "To be determined based on initial launch feedback.", color=MUTED)
    doc.add_paragraph()

    # ═══════════════════════════════════════════
    #  SECTION 6: THE COST TO MAKE IT
    # ═══════════════════════════════════════════

    _heading(doc, "6. Estimated Cost")

    cost_data = report.get("cost_estimate")
    if cost_data and isinstance(cost_data, dict):
        for k, v in cost_data.items():
            label = k.replace("_", " ").title()
            p = doc.add_paragraph()
            r1 = p.add_run(f"{label}: ")
            r1.bold = True
            r1.font.size = Pt(11)
            r1.font.color.rgb = DARK
            r2 = p.add_run(str(v))
            r2.font.size = Pt(11)
    elif cost_data and isinstance(cost_data, str):
        _para(doc, cost_data)
    else:
        tech_stack = arch_c.get("tech_stack", [])
        infra = _extract(arch_c, "infrastructure", "hosting", "deployment")

        _para(doc, "Cost breakdown based on architecture decisions:", bold=True)
        doc.add_paragraph()

        cost_items = []
        if tech_stack and isinstance(tech_stack, list):
            for item in tech_stack:
                if isinstance(item, dict):
                    choice = item.get("choice", item.get("name", ""))
                    cat = item.get("category", "")
                    if choice:
                        cost_items.append(f"{cat}: {choice} (open-source / free tier)")
        if infra:
            cost_items.append(f"Infrastructure: {infra}")

        if cost_items:
            _bullet_list(doc, cost_items)

        _para(doc, "Note: Detailed cost estimation depends on scale, hosting choices, and team size. "
                    "Most MVP builds using open-source stacks can launch for $0-500/month in infrastructure costs.",
              size=10, italic=True, color=MUTED)

    # ═══════════════════════════════════════════
    #  APPENDIX: DETAILED AGENT OUTPUTS
    # ═══════════════════════════════════════════

    sections = [
        ("ceo", "Executive Summary"),
        ("business_analyst", "Requirements Analysis"),
        ("researcher", "Market Research"),
        ("architect", "Technical Architecture"),
        ("engineer", "Implementation Plan"),
    ]

    for role, section_title in sections:
        output = _get_output_by_role(outputs, role)
        if not output:
            continue

        doc.add_page_break()
        _heading(doc, f"Appendix: {section_title}")

        content = output.get("content", {})
        if isinstance(content, dict):
            display = {k: v for k, v in content.items() if not k.startswith("_")}
            if "raw_response" in display and len(display) <= 2:
                _para(doc, str(display["raw_response"]))
            else:
                _add_structured(doc, display)
        else:
            _para(doc, str(content))

        review_key = f"peer_review_{role}"
        review_raw = memory.get(review_key)
        if review_raw:
            try:
                review = json.loads(review_raw)
                doc.add_paragraph()
                h = doc.add_heading(f"Peer Review by {review.get('reviewer_label', 'Teammate')}", level=2)
                for run in h.runs:
                    run.font.color.rgb = WARNING

                if review.get("team_note"):
                    _para(doc, f'"{review["team_note"]}"', size=10, italic=True, color=MUTED)

                if review.get("overall_assessment"):
                    _para(doc, review["overall_assessment"])

                for label, key, color in [
                    ("Strengths", "strengths", SUCCESS),
                    ("Concerns", "concerns", WARNING),
                    ("Suggestions", "suggestions", ACCENT),
                ]:
                    items = review.get(key, [])
                    if items:
                        h = doc.add_heading(label, level=3)
                        for run in h.runs:
                            run.font.color.rgb = color
                        _bullet_list(doc, items)
            except (json.JSONDecodeError, TypeError):
                pass

    # ═══════════════════════════════════════════
    #  APPENDIX: RESEARCH SOURCES
    # ═══════════════════════════════════════════

    research_raw = memory.get("research_raw_data")
    if research_raw:
        doc.add_page_break()
        _heading(doc, "Appendix: Research Sources")
        _para(doc, "Sources discovered during automated web research.", size=10, color=MUTED)

        try:
            raw_results = json.loads(research_raw)
            for i, res in enumerate(raw_results, 1):
                doc.add_paragraph()
                p = doc.add_paragraph()
                r = p.add_run(f"{i}. {res.get('title', 'Untitled')}")
                r.bold = True
                r.font.size = Pt(10)

                p2 = doc.add_paragraph()
                r2 = p2.add_run(f"Source: {res.get('source', 'unknown')} | URL: {res.get('url', 'N/A')}")
                r2.font.size = Pt(9)
                r2.font.color.rgb = ACCENT

                if res.get("snippet"):
                    _para(doc, res["snippet"], size=9, color=MUTED)
        except (json.JSONDecodeError, TypeError):
            _para(doc, str(research_raw))

    file_path = os.path.join(GENERATED_DIR, f"{project_id}_report.docx")
    doc.save(file_path)
    return file_path


def get_docx_path(project_id: str) -> str | None:
    path = os.path.join(GENERATED_DIR, f"{project_id}_report.docx")
    return path if os.path.exists(path) else None
